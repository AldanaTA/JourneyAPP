"""
Reusable helper utilities for Journey import agents.

This module intentionally contains generic functionality that can be reused by
other agents, such as file text extraction and character-count batching.
"""

from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor, as_completed
import json
import re
from pathlib import Path
from typing import Any, Callable, Sequence, TypeVar, cast

T = TypeVar("T")
R = TypeVar("R")


def normalize_text(text: str) -> str:
    """Normalize line endings and collapse excessive blank lines."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_txt(path: Path) -> str:
    """Read plain text from a .txt file."""
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    """Extract text from a .docx file, including basic table contents."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency. Install with: pip install python-docx") from exc

    doc = Document(path)
    chunks: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    # Also read tables, because imported content may be formatted as tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def read_pdf(path: Path) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Missing dependency. Install with: pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []

    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"\n--- PAGE {i} ---\n{page_text}")

    return "\n".join(pages)


def extract_text(path: Path | str) -> str:
    """Extract text from supported file types: .txt, .docx, and .pdf."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return read_txt(path)

    if suffix == ".docx":
        return read_docx(path)

    if suffix == ".pdf":
        return read_pdf(path)

    raise ValueError(f"Unsupported file type: {suffix}. Use .txt, .docx, or .pdf.")


def batch_by_char_count(
    items: Sequence[T],
    max_chars: int,
    text_getter: Callable[[T], str],
) -> list[list[T]]:
    """
    Split items into batches using the character length of each item's text.

    Oversized individual items are placed into their own batch rather than
    being dropped or split.
    """
    if max_chars <= 0:
        raise ValueError("max_chars must be greater than 0.")

    batches: list[list[T]] = []
    current: list[T] = []
    current_size = 0

    for item in items:
        item_size = len(text_getter(item))

        if current and current_size + item_size > max_chars:
            batches.append(current)
            current = []
            current_size = 0

        current.append(item)
        current_size += item_size

    if current:
        batches.append(current)

    return batches


def run_threaded_batches(
    batches: Sequence[Sequence[T]],
    worker: Callable[[list[T]], R],
    max_workers: int = 4,
) -> list[R]:
    """
    Run a batch worker concurrently while preserving the original batch order.

    This is useful for import agents that send independent batches to an API.
    Each batch is converted to a list before being passed to the worker so the
    worker can safely iterate over it multiple times.
    """
    if max_workers <= 0:
        raise ValueError("max_workers must be greater than 0.")

    if not batches:
        return []

    worker_count = min(max_workers, len(batches))
    ordered_results: list[Any] = [None] * len(batches)

    with ThreadPoolExecutor(max_workers=worker_count) as executor:
        future_to_index = {
            executor.submit(worker, list(batch)): index
            for index, batch in enumerate(batches)
        }

        for future in as_completed(future_to_index):
            index = future_to_index[future]
            ordered_results[index] = future.result()

    return cast(list[R], ordered_results)


def write_json_file(data: Any, output_path: Path | str) -> None:
    """Write data to a UTF-8 JSON file with readable indentation."""
    output_path = Path(output_path)
    output_path.write_text(
        json.dumps(data, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
