"""
power_import_agent_updated.py

Updated supervised import agent for your `powers` Postgres table.

This version is hardened for documents shaped like your sample files:

Power Name
LVL:
Type:
Cost:
Material Components:
Components:
Range:
Area:
Duration:
Effect:
Empower:
LVL UP:

Major improvements:
- Splits the document into power blocks before calling the model.
- Ignores intro/rules text and headers like EP Powers / MP Powers.
- Explicitly maps Components: V,O,S,D,C into database booleans.
- Parses costs into hp_cost/mp_cost/ep_cost/tp_cost.
- Flags variable costs like Cost: *MP 2TP as invalid.
- Keeps alternate cast-time notes by placing them in the effect text.
- Handles both "LVL UP:" and mistaken final "LVL:" lines as level-up text.

Install:
    pip install -r requirements.txt

Environment:
    OPENAI_API_KEY=your_api_key
    DATABASE_URL=postgresql://user:password@localhost:5432/dbname
    OPENAI_MODEL=gpt-4.1-mini

Preview:
    python power_import_agent_updated.py "./Powers(Journey 100).txt" --out import_preview.json

Commit:
    python power_import_agent_updated.py "./Powers(Journey 100).txt" --commit
"""

from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

from dotenv import load_dotenv
from openai import OpenAI
from pydantic import BaseModel, ConfigDict, Field, ValidationError, field_validator


# ----------------------------
# Your DB enum
# ----------------------------

PowerType = Literal["Destruction", "Support", "Sabotage", "Utility"]


# ----------------------------
# Pydantic validation models
# ----------------------------

class PowerRow(BaseModel):
    model_config = ConfigDict(extra="forbid")

    name: str = Field(min_length=1)
    lvl: int = Field(ge=1)
    type: PowerType

    tp_cost: int = Field(ge=0)
    hp_cost: int | None = Field(default=None, ge=0)
    mp_cost: int | None = Field(default=None, ge=0)
    ep_cost: int | None = Field(default=None, ge=0)

    material_components: str = "None"
    verbal: bool
    sight: bool
    somatic: bool
    is_distinct: bool
    concentration: bool

    range: str = "5"
    area: str = "0"
    duration: str = "Instantaneous"

    effect: str = Field(min_length=1)
    empower_effect: str | None = None
    lvl_up_effect: str | None = None

    @field_validator(
        "name",
        "material_components",
        "range",
        "area",
        "duration",
        "effect",
        mode="before",
    )
    @classmethod
    def clean_required_strings(cls, value: Any) -> str:
        if value is None:
            return ""
        return str(value).strip()

    @field_validator("empower_effect", "lvl_up_effect", mode="before")
    @classmethod
    def clean_optional_strings(cls, value: Any) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        if text == "" or text.lower() in {"none", "n/a", "null", "-"}:
            return None
        return text

    @field_validator("material_components", mode="after")
    @classmethod
    def default_material_components(cls, value: str) -> str:
        if value.strip() in {"", "-"}:
            return "None"
        return value

    @field_validator("range", mode="after")
    @classmethod
    def default_range(cls, value: str) -> str:
        if value.strip() in {"", "None", "none", "-"}:
            return "5"
        return value

    @field_validator("area", mode="after")
    @classmethod
    def default_area(cls, value: str) -> str:
        if value.strip() in {"", "None", "none", "-"}:
            return "0"
        return value

    @field_validator("duration", mode="after")
    @classmethod
    def default_duration(cls, value: str) -> str:
        if value.strip() in {"", "None", "none", "-"}:
            return "Instantaneous"
        return value


class InvalidPower(BaseModel):
    model_config = ConfigDict(extra="forbid")

    source_text: str
    errors: list[str]


class AgentImportResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    valid_powers: list[PowerRow]
    invalid_powers: list[InvalidPower]


# ----------------------------
# Power block splitting
# ----------------------------

@dataclass(frozen=True)
class PowerBlock:
    index: int
    name: str
    text: str


IGNORED_HEADER_LINES = {
    "powers",
    "casting powers",
    "empowering powers",
    "destructive powers",
    "destruction powers",
    "support powers",
    "sabotage powers",
    "utility powers",
    "ep powers",
    "mp powers",
}


REQUIRED_POWER_LABELS = [
    "LVL:",
    "Type:",
    "Cost:",
    "Material Components:",
    "Components:",
    "Range:",
    "Area:",
    "Duration:",
    "Effect:",
]


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def is_possible_power_name(line: str) -> bool:
    cleaned = line.strip()

    if not cleaned:
        return False

    if ":" in cleaned:
        return False

    if cleaned.lower() in IGNORED_HEADER_LINES:
        return False

    # Avoid very long rules paragraphs being treated as names.
    if len(cleaned) > 80:
        return False

    # Most power names are title-like, but allow names such as "Weaken Damage Type".
    if not re.search(r"[A-Za-z]", cleaned):
        return False

    return True


def has_required_power_labels(block_text: str) -> bool:
    return all(label in block_text for label in REQUIRED_POWER_LABELS)


def split_power_blocks(raw_text: str) -> list[PowerBlock]:
    """
    Finds power blocks by looking for:

        Some Power Name
        LVL: ...

    Then captures until the next line that looks like another power name
    followed by LVL:.

    This intentionally ignores intro/rules text and category headers.
    """
    text = normalize_text(raw_text)
    lines = [line.rstrip() for line in text.split("\n")]

    start_indexes: list[int] = []

    for i in range(len(lines) - 1):
        current_line = lines[i].strip()
        next_line = lines[i + 1].strip()

        if is_possible_power_name(current_line) and next_line.startswith("LVL:"):
            start_indexes.append(i)

    blocks: list[PowerBlock] = []

    for block_number, start in enumerate(start_indexes, start=1):
        end = start_indexes[block_number] if block_number < len(start_indexes) else len(lines)
        block_lines = lines[start:end]
        block_text = "\n".join(block_lines).strip()
        name = block_lines[0].strip()

        if has_required_power_labels(block_text):
            blocks.append(PowerBlock(index=block_number, name=name, text=block_text))

    return blocks


# ----------------------------
# Pre-validation
# ----------------------------

def get_label_value(block_text: str, label: str) -> str | None:
    pattern = rf"(?im)^{re.escape(label)}\s*(.*)$"
    match = re.search(pattern, block_text)
    if not match:
        return None
    return match.group(1).strip()


def prevalidate_power_block(block: PowerBlock) -> list[str]:
    """
    Catches cases that your DB schema cannot represent safely.

    We only reject things that would clearly break NOT NULL / INT columns
    or cause major data loss.
    """
    errors: list[str] = []

    lvl_raw = get_label_value(block.text, "LVL:")
    cost_raw = get_label_value(block.text, "Cost:")
    type_raw = get_label_value(block.text, "Type:")

    if not lvl_raw:
        errors.append("Missing LVL.")
    elif not re.fullmatch(r"\d+", lvl_raw):
        errors.append(f"LVL must be an integer, got: {lvl_raw!r}.")

    if not type_raw:
        errors.append("Missing Type.")
    elif type_raw.strip() not in {"Destruction", "Support", "Sabotage", "Utility"}:
        errors.append(f"Invalid Type: {type_raw!r}.")

    if not cost_raw:
        errors.append("Missing Cost.")
    else:
        if "*" in cost_raw:
            errors.append(
                f"Variable cost cannot fit mp_cost/hp_cost/ep_cost integer columns: {cost_raw!r}."
            )

        if not re.search(r"(?i)\b\d+\s*TP\b", cost_raw):
            errors.append(f"Missing numeric TP cost in Cost: {cost_raw!r}.")

    return errors


def remove_invalid_blocks(blocks: list[PowerBlock]) -> tuple[list[PowerBlock], list[InvalidPower]]:
    valid_for_agent: list[PowerBlock] = []
    invalid_powers: list[InvalidPower] = []

    for block in blocks:
        errors = prevalidate_power_block(block)

        if errors:
            invalid_powers.append(InvalidPower(source_text=block.text, errors=errors))
        else:
            valid_for_agent.append(block)

    return valid_for_agent, invalid_powers


# ----------------------------
# Structured output schema
# ----------------------------

POWER_IMPORT_JSON_SCHEMA: dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "valid_powers": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "name": {"type": "string"},
                    "lvl": {"type": "integer"},
                    "type": {
                        "type": "string",
                        "enum": ["Destruction", "Support", "Sabotage", "Utility"],
                    },
                    "tp_cost": {"type": "integer"},
                    "hp_cost": {"type": ["integer", "null"]},
                    "mp_cost": {"type": ["integer", "null"]},
                    "ep_cost": {"type": ["integer", "null"]},
                    "material_components": {"type": "string"},
                    "verbal": {"type": "boolean"},
                    "sight": {"type": "boolean"},
                    "somatic": {"type": "boolean"},
                    "is_distinct": {"type": "boolean"},
                    "concentration": {"type": "boolean"},
                    "range": {"type": "string"},
                    "area": {"type": "string"},
                    "duration": {"type": "string"},
                    "effect": {"type": "string"},
                    "empower_effect": {"type": ["string", "null"]},
                    "lvl_up_effect": {"type": ["string", "null"]},
                },
                "required": [
                    "name",
                    "lvl",
                    "type",
                    "tp_cost",
                    "hp_cost",
                    "mp_cost",
                    "ep_cost",
                    "material_components",
                    "verbal",
                    "sight",
                    "somatic",
                    "is_distinct",
                    "concentration",
                    "range",
                    "area",
                    "duration",
                    "effect",
                    "empower_effect",
                    "lvl_up_effect",
                ],
            },
        },
        "invalid_powers": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "source_text": {"type": "string"},
                    "errors": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                },
                "required": ["source_text", "errors"],
            },
        },
    },
    "required": ["valid_powers", "invalid_powers"],
}


AGENT_INSTRUCTIONS = """
You are a supervised data import agent for a tabletop power database.

You will receive pre-split power blocks. Each block should represent exactly one power.
Do not extract intro/rules text, category headers, or section headers.

Return only JSON matching the provided schema.

Database columns:
- name TEXT NOT NULL
- lvl INT NOT NULL
- type enum: Destruction, Support, Sabotage, Utility
- tp_cost INT NOT NULL
- hp_cost INT nullable
- mp_cost INT nullable
- ep_cost INT nullable
- material_components TEXT DEFAULT 'None'
- verbal BOOLEAN NOT NULL
- sight BOOLEAN NOT NULL
- somatic BOOLEAN NOT NULL
- is_distinct BOOLEAN NOT NULL
- concentration BOOLEAN NOT NULL
- range TEXT DEFAULT '5'
- area TEXT DEFAULT '0'
- duration TEXT DEFAULT 'Instantaneous'
- effect TEXT NOT NULL
- empower_effect TEXT nullable
- lvl_up_effect TEXT nullable

Cost parsing rules:
- Parse "Cost:" into tp_cost, hp_cost, mp_cost, and ep_cost.
- Example: "Cost: 20EP 15HP 4TP" means ep_cost=20, hp_cost=15, tp_cost=4, mp_cost=null.
- Example: "Cost: 55MP 3TP" means mp_cost=55, tp_cost=3, hp_cost=null, ep_cost=null.
- If a resource does not appear, use null for that resource.
- TP must be a number.
- If Cost contains an alternate cast time, such as "12MP 4TP or 1 minute to cast", parse the numeric costs normally and add this sentence to the beginning of effect: "Alternate casting: 1 minute to cast."
- If a block has variable cost like "*MP", place it in invalid_powers.

Component parsing rules:
The source field "Components:" maps to database booleans:
- V means verbal = true
- O means sight = true
- S means somatic = true
- D means is_distinct = true
- C means concentration = true
- None means all five booleans are false
- Missing letters are false
- Ignore commas and whitespace

Field mapping rules:
- "Material Components:" maps to material_components.
- Empty material components, "-", or "None" should become "None".
- "Range: None" should become "5" because the DB default is "5".
- "Area: None" should become "0" because the DB default is "0".
- "Duration: None" should become "Instantaneous" because the DB default is "Instantaneous".
- "Empower:" maps to empower_effect. If it says "None", use null.
- "LVL UP:" maps to lvl_up_effect. If it says "None", use null.
- Some source blocks mistakenly use a final "LVL:" line where they clearly mean "LVL UP:". If a second LVL-like field appears near the end after Empower, treat it as lvl_up_effect.
- Preserve the Effect/Empower/LVL UP wording as much as possible.

Invalid rules:
Put a block into invalid_powers when:
- name is missing
- lvl is missing, "None", or not an integer
- type is not one of Destruction, Support, Sabotage, Utility
- Cost is missing
- Cost has variable resource values like "*MP"
- tp_cost cannot be parsed
- effect is missing
"""


# ----------------------------
# File text extraction
# ----------------------------

def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency. Install with: pip install python-docx") from exc

    doc = Document(path)
    chunks: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    # Also read tables, because powers may later be formatted as tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Missing dependency. Install with: pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []

    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"\n--- PAGE {i} ---\n{page_text}")

    return "\n".join(pages)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return read_txt(path)

    if suffix == ".docx":
        return read_docx(path)

    if suffix == ".pdf":
        return read_pdf(path)

    raise ValueError(f"Unsupported file type: {suffix}. Use .txt, .docx, or .pdf.")


# ----------------------------
# Agent batching
# ----------------------------

def make_block_batch_text(blocks: list[PowerBlock]) -> str:
    parts: list[str] = []

    for block in blocks:
        parts.append(f"--- POWER BLOCK {block.index}: {block.name} ---\n{block.text}")

    return "\n\n".join(parts)


def batch_blocks(blocks: list[PowerBlock], max_chars: int) -> list[list[PowerBlock]]:
    batches: list[list[PowerBlock]] = []
    current: list[PowerBlock] = []
    current_size = 0

    for block in blocks:
        block_size = len(block.text)

        if current and current_size + block_size > max_chars:
            batches.append(current)
            current = []
            current_size = 0

        current.append(block)
        current_size += block_size

    if current:
        batches.append(current)

    return batches


def extract_powers_batch_with_agent(blocks: list[PowerBlock], model: str) -> AgentImportResult:
    client = OpenAI()

    response = client.responses.create(
        model=model,
        instructions=AGENT_INSTRUCTIONS,
        input=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": (
                            "Extract the following pre-split power blocks into database rows.\n\n"
                            f"{make_block_batch_text(blocks)}"
                        ),
                    }
                ],
            }
        ],
        text={
            "format": {
                "type": "json_schema",
                "name": "power_import_result",
                "strict": True,
                "schema": POWER_IMPORT_JSON_SCHEMA,
            }
        },
    )

    raw_json = response.output_text
    parsed = json.loads(raw_json)

    return AgentImportResult.model_validate(parsed)


def merge_results(results: list[AgentImportResult]) -> AgentImportResult:
    valid_powers: list[PowerRow] = []
    invalid_powers: list[InvalidPower] = []

    for result in results:
        valid_powers.extend(result.valid_powers)
        invalid_powers.extend(result.invalid_powers)

    return AgentImportResult(valid_powers=valid_powers, invalid_powers=invalid_powers)


def extract_powers_with_agent(
    raw_text: str,
    model: str,
    max_batch_chars: int,
) -> tuple[AgentImportResult, dict[str, Any]]:
    all_blocks = split_power_blocks(raw_text)
    blocks_for_agent, pre_invalid = remove_invalid_blocks(all_blocks)

    metadata: dict[str, Any] = {
        "power_blocks_found": len(all_blocks),
        "pre_validation_invalid": len(pre_invalid),
        "sent_to_agent": len(blocks_for_agent),
        "batches": 0,
    }

    if not all_blocks:
        return (
            AgentImportResult(
                valid_powers=[],
                invalid_powers=[
                    InvalidPower(
                        source_text=raw_text[:4000],
                        errors=[
                            "No power blocks found. Expected a power name line followed by 'LVL:'."
                        ],
                    )
                ],
            ),
            metadata,
        )

    batches = batch_blocks(blocks_for_agent, max_chars=max_batch_chars)
    metadata["batches"] = len(batches)

    batch_results: list[AgentImportResult] = []

    if pre_invalid:
        batch_results.append(AgentImportResult(valid_powers=[], invalid_powers=pre_invalid))

    for batch in batches:
        batch_results.append(extract_powers_batch_with_agent(batch, model=model))

    return merge_results(batch_results), metadata


# ----------------------------
# Import preview output
# ----------------------------

def write_preview(result: AgentImportResult, metadata: dict[str, Any], output_path: Path) -> None:
    preview = {
        "metadata": metadata,
        "valid_powers": [power.model_dump() for power in result.valid_powers],
        "invalid_powers": [power.model_dump() for power in result.invalid_powers],
    }

    output_path.write_text(
        json.dumps(preview, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


# ----------------------------
# Database insert
# ----------------------------

INSERT_POWER_SQL = """
INSERT INTO powers (
    name,
    lvl,
    type,
    tp_cost,
    hp_cost,
    mp_cost,
    ep_cost,
    material_components,
    verbal,
    sight,
    somatic,
    is_distinct,
    concentration,
    range,
    area,
    duration,
    effect,
    empower_effect,
    lvl_up_effect
)
VALUES (
    %(name)s,
    %(lvl)s,
    %(type)s,
    %(tp_cost)s,
    %(hp_cost)s,
    %(mp_cost)s,
    %(ep_cost)s,
    %(material_components)s,
    %(verbal)s,
    %(sight)s,
    %(somatic)s,
    %(is_distinct)s,
    %(concentration)s,
    %(range)s,
    %(area)s,
    %(duration)s,
    %(effect)s,
    %(empower_effect)s,
    %(lvl_up_effect)s
)
RETURNING id;
"""


def insert_valid_powers(result: AgentImportResult, database_url: str) -> list[str]:
    try:
        import psycopg
    except ImportError as exc:
        raise RuntimeError("Missing dependency. Install with: pip install psycopg[binary]") from exc

    inserted_ids: list[str] = []

    with psycopg.connect(database_url) as conn:
        with conn.cursor() as cur:
            for power in result.valid_powers:
                cur.execute(INSERT_POWER_SQL, power.model_dump())
                inserted_id = cur.fetchone()[0]
                inserted_ids.append(str(inserted_id))

        conn.commit()

    return inserted_ids


# ----------------------------
# Main CLI
# ----------------------------

def main() -> None:
    load_dotenv()

    parser = argparse.ArgumentParser(
        description="Import powers from a TXT, DOCX, or PDF file into a reviewable JSON preview."
    )
    parser.add_argument("file", help="Path to .txt, .docx, or .pdf file.")
    parser.add_argument(
        "--model",
        default=os.getenv("OPENAI_MODEL", "gpt-4.1-mini"),
        help="OpenAI model to use. Can also be set with OPENAI_MODEL.",
    )
    parser.add_argument(
        "--out",
        default="power_import_preview.json",
        help="Where to save the reviewable JSON preview.",
    )
    parser.add_argument(
        "--max-batch-chars",
        type=int,
        default=24000,
        help="Approximate max characters sent to the model per batch.",
    )
    parser.add_argument(
        "--commit",
        action="store_true",
        help="Insert valid powers into Postgres after preview generation.",
    )

    args = parser.parse_args()

    source_path = Path(args.file)
    output_path = Path(args.out)

    if not source_path.exists():
        raise FileNotFoundError(source_path)

    if not os.getenv("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY is missing. Put it in your environment or .env file.")

    print(f"Reading: {source_path}")
    raw_text = extract_text(source_path)

    if not raw_text.strip():
        raise RuntimeError("No text was extracted from the file.")

    print("Splitting document into power blocks and extracting with agent...")

    try:
        result, metadata = extract_powers_with_agent(
            raw_text,
            model=args.model,
            max_batch_chars=args.max_batch_chars,
        )
    except ValidationError as exc:
        print("The agent returned JSON, but it failed local validation.")
        raise exc

    write_preview(result, metadata, output_path)

    print()
    print("Import preview created.")
    print(f"Power blocks found:       {metadata['power_blocks_found']}")
    print(f"Pre-validation invalid:   {metadata['pre_validation_invalid']}")
    print(f"Sent to agent:            {metadata['sent_to_agent']}")
    print(f"Batches:                  {metadata['batches']}")
    print(f"Valid powers:             {len(result.valid_powers)}")
    print(f"Invalid powers:           {len(result.invalid_powers)}")
    print(f"Preview file:             {output_path}")

    if result.invalid_powers:
        print()
        print("Invalid powers found. Review the preview before committing.")

    if args.commit:
        database_url = os.getenv("DATABASE_URL")
        if not database_url:
            raise RuntimeError("DATABASE_URL is missing. Required when using --commit.")

        print()
        print("Committing valid powers to database...")
        inserted_ids = insert_valid_powers(result, database_url)
        print(f"Inserted {len(inserted_ids)} powers.")


if __name__ == "__main__":
    main()
